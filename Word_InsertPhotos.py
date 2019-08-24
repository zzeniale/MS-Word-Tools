"""

insert photos into word table

"""
import os
import docx
import tkinter as tk
from tkinter import filedialog, simpledialog
import ctypes

application_window = tk.Tk()

# Ask the user to select a folder
imagefolder = filedialog.askdirectory(parent=application_window,
                              initialdir=os.getcwd(),
                              title="Please select the image folder.")
suffix = simpledialog.askstring("Input", "What is the image suffix? E.g. SIS01-T, LNR02-N",
                                parent=application_window)
application_window.destroy()


#create word document and set style
doc = docx.Document()
font = doc.styles['Normal'].font
font.name = 'Arial'
font.size = docx.shared.Pt(9)

#create table in word document
table = doc.add_table(rows = 1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Colony number"
hdr_cells[1].text = "Latest monitoring period"

#insert image into table rows
imgnames = os.listdir(imagefolder)
for i in range(len(imgnames)):
    imgnames[i]=int(str(imgnames[i])[8:-4])    
imgnames = sorted(imgnames)
for image in imgnames: 
    imgpath = os.path.join(imagefolder,suffix+'-'+ str(image)+'.jpg')
    newrow = table.add_row().cells
    paragraph = newrow[1].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(imgpath,width = docx.shared.Cm(2.79))
    newrow[0].text = str(image)

#table width and paragraph spacing
for cell in table.columns[0].cells:
    cell.width = docx.shared.Cm(1.5)
for cell in table.columns[1].cells:
    cell.width = docx.shared.Cm(2.9)
for row in table.rows:
    for cell in row.cells:
        paragraph = cell.paragraphs[0]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = docx.shared.Pt(0)
        paragraph_format.space_after = docx.shared.Pt(0)

#save document
doc.save(os.path.join(imagefolder, "image_table.docx"))

ctypes.windll.user32.MessageBoxW(0, "All done!", "Message", 1)
